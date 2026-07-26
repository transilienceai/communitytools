#!/usr/bin/env python3
"""Parse a finished/foreign report into the canonical report_data finding schema.

So every downstream flow (merge / retest / CERT-In / consolidation) can bootstrap
from a finished report instead of hand re-keying findings. This tool owns the
DETERMINISTIC extraction; fuzzy free-form-narrative interpretation stays a
caller/LLM step.

Supported sources (auto-detected by extension, then by magic-byte sniff):
  * json  -- a native report_data.json: validate shape + copy findings (passthrough).
  * xlsx  -- openpyxl; reads the "Findings Register" sheet as the faithful INVERSE
             of tools/report_xlsx_build.py (its columns -> finding fields).
  * md    -- a findings markdown: BOTH a `| ID | Title | Severity | ... |` table
             AND `## <id> - <title>` sectioned findings (Severity:/CVSS:/CWE: lines).
  * pdf   -- best-effort text via `pdftotext -layout` (subprocess) or pdfminer.
  * docx  -- best-effort via python-docx, else the raw word/document.xml text.

Output: {engagement:{...}, findings:[...], ingest:{source, format, warnings:[...]}}

NON-NULL GUARDS: after extraction each finding is checked and a warning is emitted
(into ingest.warnings) when cwe is None or cvss_vector is None -- the lossy-PDF
'cwe:None' regression. Values are NEVER invented.

Exit: 0 clean / 3 a format's optional dependency/tool is missing (partial or empty
output + warning) / 2 usage or unreadable source.

STDLIB-ONLY at import time -- per-format deps (openpyxl / python-docx /
pdftotext|pdfminer) are imported lazily; their absence degrades to a warning +
empty-findings shell (exit 3), never a crash.

REUSABLE CONTENT: this tool carries NO client name, host, or engagement path.

Usage:
    python3 tools/report_ingest.py <source> [--format auto|json|xlsx|md|pdf|docx]
                                            [-o OUT]
"""
import argparse
import html
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile

HERE = os.path.dirname(os.path.abspath(__file__))
if HERE not in sys.path:
    sys.path.insert(0, HERE)

# report_data_shape is stdlib-only and imports cleanly with no reportlab; used to
# lint a json passthrough. Optional -- ingest still works without it.
try:
    from report_data_shape import require_report_data_shape
except Exception:  # pragma: no cover - defensive
    require_report_data_shape = None

EXIT_OK = 0
EXIT_USAGE = 2
EXIT_MISSING_DEP = 3

# Placeholders the report builders emit for an absent value (report_xlsx_build).
_EMPTY = {"", "-", "—", "n/a", "na", "none"}


class MissingDependency(Exception):
    """A format's optional dependency/tool is unavailable (-> exit 3)."""


class UnreadableSource(Exception):
    """The source cannot be read or is not the claimed format (-> exit 2)."""


# --- header/label vocabulary --------------------------------------------------
# canonical finding field -> the header/label spellings that map to it. The xlsx
# spellings ("Affected Hosts", "CWE / Classification", "CVE(s)") mirror
# report_xlsx_build.build_register()'s header row exactly, so the xlsx read is its
# faithful inverse; the extra spellings absorb foreign markdown/docx reports.
HEADER_ALIASES = {
    "id": {"id", "finding id", "ref id"},
    "title": {"finding", "title", "name", "vulnerability", "issue", "summary"},
    "severity": {"severity", "sev", "risk", "risk rating"},
    "cvss_score": {"cvss", "cvss score", "score", "cvss base score", "base score"},
    "cvss_vector": {"cvss vector", "vector", "cvss v3 vector", "cvss v4 vector"},
    "affected": {"affected hosts", "affected", "affected host", "assets", "asset",
                 "hosts", "host", "target", "targets", "endpoint", "endpoints",
                 "url", "location", "affected assets"},
    "cwe": {"cwe / classification", "cwe", "classification", "cwe/owasp",
            "weakness", "cwe id"},
    "owasp": {"owasp", "owasp category"},
    "cves": {"cve(s)", "cves", "cve", "cve ids", "cve id"},
    "status": {"status"},
    "references": {"references", "reference", "refs"},
    "description": {"description", "desc", "details", "detail"},
    "impact": {"impact"},
    "recommendation": {"recommendation", "remediation", "fix", "mitigation"},
}
_HEADER_TO_FIELD = {a: f for f, aliases in HEADER_ALIASES.items() for a in aliases}

# Summary-sheet Scope & Method label -> engagement key (inverse of build_summary).
SUMMARY_LABELS = {
    "client / target": "target", "target": "target", "client": "target",
    "sector": "sector", "method": "method", "report id": "report_id",
    "date": "date", "classification": "classification",
}

_CWE_RE = re.compile(r"CWE-\d+", re.I)
_OWASP_RE = re.compile(r"\b(?:A\d{1,2}:\d{4}|API\d{1,2}:\d{4}|M\d{1,2}:\d{4})\b", re.I)
_CVE_RE = re.compile(r"CVE-\d{4}-\d{4,7}", re.I)
_SCORE_RE = re.compile(r"\d+(?:\.\d+)?")


def _canon_header(h):
    """Map a header/label spelling to a canonical finding field, or '' if unknown."""
    key = re.sub(r"\s+", " ", str(h or "").strip().lower()).strip(" *:")
    return _HEADER_TO_FIELD.get(key, "")


def _is_empty(value):
    return value is None or str(value).strip().lower() in _EMPTY


def norm_severity(s):
    """Canonicalize to the schema band; 'Informational' -> 'Info'. Unknown bands
    pass through title-cased (never invented) for a caller to reconcile."""
    if s is None:
        return None
    t = str(s).strip()
    if not t:
        return None
    t = t.title()
    return "Info" if t in ("Info", "Informational") else t


def parse_cvss_score(v):
    if _is_empty(v):
        return None
    m = _SCORE_RE.search(str(v))
    return float(m.group(0)) if m else None


def parse_cvss_vector(v):
    if _is_empty(v):
        return None
    s = str(v).strip()
    return s if (s.upper().startswith("CVSS:") or "AV:" in s.upper()) else None


def parse_cwe(v):
    if _is_empty(v):
        return None
    toks = [t.upper() for t in _CWE_RE.findall(str(v))]
    seen, out = set(), []
    for t in toks:
        if t not in seen:
            seen.add(t)
            out.append(t)
    return "; ".join(out) if out else None


def parse_owasp(v):
    if _is_empty(v):
        return None
    toks = [t.upper() for t in _OWASP_RE.findall(str(v))]
    seen, out = set(), []
    for t in toks:
        if t not in seen:
            seen.add(t)
            out.append(t)
    return "; ".join(out) if out else None


def _split(value, comma=False):
    """Split a multi-value cell on newline / <br> / ';' (and ',' when comma=True),
    dropping placeholders and bullet glyphs."""
    if value is None:
        return []
    pat = r"<br\s*/?>|[\n;" + ("," if comma else "") + r"]"
    out = []
    for chunk in re.split(pat, str(value), flags=re.I):
        c = chunk.strip().strip("•").strip()
        if c and not _is_empty(c):
            out.append(c)
    return out


def parse_cves(v):
    out = []
    for tok in _split(v, comma=True):
        m = _CVE_RE.search(tok)
        if m:
            out.append({"id": m.group(0).upper()})
    return out


def _build_finding(values, *, affected_comma):
    """Assemble one canonical finding from a {canonical_field: raw_value} dict.
    Only present, non-placeholder fields are emitted; nothing is invented."""
    f = {"id": str(values.get("id") or "").strip(),
         "title": str(values.get("title") or "").strip()}
    sev = norm_severity(values.get("severity"))
    if sev:
        f["severity"] = sev
    score = parse_cvss_score(values.get("cvss_score"))
    if score is not None:
        f["cvss_score"] = score
    vec = parse_cvss_vector(values.get("cvss_vector"))
    if vec:
        f["cvss_vector"] = vec
    cwe_src = values.get("cwe")
    cwe = parse_cwe(cwe_src)
    if cwe:
        f["cwe"] = cwe
    owasp = parse_owasp(values.get("owasp")) or parse_owasp(cwe_src)
    if owasp:
        f["owasp"] = owasp
    affected = _split(values.get("affected"), comma=affected_comma)
    if affected:
        f["affected"] = affected
    cves = parse_cves(values.get("cves"))
    if cves:
        f["cves"] = cves
    refs = _split(values.get("references"), comma=True)
    if refs:
        f["references"] = refs
    for k in ("description", "impact", "recommendation"):
        val = values.get(k)
        if val is not None and not _is_empty(val):
            f[k] = str(val).strip()
    if "needs live" in str(values.get("status") or "").lower():
        f["needs_live_confirmation"] = True
    return f


def _findings_from_grid(rows, affected_comma):
    """Locate a findings header row in a grid (list of cell-lists) and read every
    subsequent row as a finding. Returns (findings, warning_or_None)."""
    header_idx = None
    for i, row in enumerate(rows):
        canon = [_canon_header(c) for c in row]
        if "id" in canon and ("title" in canon or "severity" in canon):
            header_idx = i
            break
    if header_idx is None:
        return [], "no findings table header (need an 'ID' column) found"
    canon = [_canon_header(c) for c in rows[header_idx]]
    out = []
    for row in rows[header_idx + 1:]:
        values = {}
        for k, c in zip(canon, row):
            if k and c is not None and str(c).strip() != "":
                values.setdefault(k, c)
        idv = str(values.get("id") or "").strip()
        if not idv or _is_empty(idv):
            continue
        out.append(_build_finding(values, affected_comma=affected_comma))
    return out, None


# --- json ---------------------------------------------------------------------
def ingest_json(path):
    try:
        with open(path, encoding="utf-8") as fh:
            data = json.load(fh)
    except (OSError, ValueError) as e:
        raise UnreadableSource("not valid JSON: %s" % e)
    if not isinstance(data, dict) or not isinstance(data.get("findings"), list):
        raise UnreadableSource(
            "JSON is not a report_data document (missing a findings[] list)")
    warnings = []
    if require_report_data_shape is not None:
        try:
            require_report_data_shape(data)
        except ValueError as e:
            warnings.append("shape warning: %s" % e)
    engagement = data.get("engagement") or {}
    findings = [dict(f) for f in data["findings"] if isinstance(f, dict)]
    return engagement, findings, warnings


# --- xlsx ---------------------------------------------------------------------
def _engagement_from_summary(wb):
    ws = next((wb[n] for n in wb.sheetnames if n.strip().lower() == "summary"), None)
    eng = {}
    if ws is None:
        return eng
    for row in ws.iter_rows(values_only=True):
        if not row or row[0] is None:
            continue
        key = SUMMARY_LABELS.get(str(row[0]).strip().lower())
        if key and len(row) > 1 and not _is_empty(row[1]):
            eng.setdefault(key, str(row[1]).strip())
    return eng


def ingest_xlsx(path):
    try:
        import openpyxl
    except ImportError:
        raise MissingDependency(
            "openpyxl not installed (pip install openpyxl); cannot read xlsx")
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception as e:
        raise UnreadableSource("cannot open xlsx: %s" % e)
    warnings = []
    # Prefer the register sheet; otherwise scan every sheet for a findings grid.
    order = ([n for n in wb.sheetnames if n.strip().lower() == "findings register"]
             + [n for n in wb.sheetnames if n.strip().lower() != "findings register"])
    findings = []
    for name in order:
        rows = [list(r) for r in wb[name].iter_rows(values_only=True)]
        fs, _ = _findings_from_grid(rows, affected_comma=False)
        if fs:
            findings = fs
            break
    if not findings:
        warnings.append("no findings register grid found in workbook")
    return _engagement_from_summary(wb), findings, warnings


# --- markdown -----------------------------------------------------------------
def _split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_sep_row(line):
    s = line.strip()
    return bool(re.fullmatch(r"[\s|:\-]+", s)) and "-" in s and "|" in s


def parse_md_table(text):
    """Parse every GitHub-flavoured findings table (header, `---` separator, rows)."""
    findings = []
    lines = text.splitlines()
    i, n = 0, len(text.splitlines())
    while i < n:
        if lines[i].strip().startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            if len(block) >= 2 and _is_sep_row(block[1]):
                grid = [_split_row(block[0])] + [_split_row(r) for r in block[2:]]
                fs, _ = _findings_from_grid(grid, affected_comma=True)
                findings.extend(fs)
        else:
            i += 1
    return findings


_SECTION_RE = re.compile(r"^\s{0,3}#{2,4}\s+(.+?)\s*$")
_ID_TITLE_RE = re.compile(r"^([A-Za-z0-9][\w.\-]*?)\s+[—–\-]\s+(.+)$")
_FIELD_LINE_RE = re.compile(
    r"^\s*[-*]?\s*\*{0,2}([A-Za-z][A-Za-z /()]*?)\*{0,2}\s*:\s*(.+?)\s*$")


def _split_id_title(heading):
    m = _ID_TITLE_RE.match(heading.strip())
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None, heading.strip()


def parse_md_sections(text):
    """Parse `## <id> - <title>` sectioned findings with Severity:/CVSS:/CWE: lines.
    A section counts as a finding only if its heading carries an id OR its body has
    a Severity: line -- so prose sections (## Executive Summary) are skipped."""
    lines = text.splitlines()
    heads = [(i, m.group(1)) for i, ln in enumerate(lines)
             for m in (_SECTION_RE.match(ln),) if m]
    findings = []
    for idx, (li, heading) in enumerate(heads):
        fid, title = _split_id_title(heading)
        end = heads[idx + 1][0] if idx + 1 < len(heads) else len(lines)
        values = {"id": fid or "", "title": title or ""}
        for bl in lines[li + 1:end]:
            m = _FIELD_LINE_RE.match(bl)
            if m:
                key = _canon_header(m.group(1))
                if key:
                    values.setdefault(key, m.group(2).strip())
        if fid is not None or "severity" in values:
            findings.append(_build_finding(values, affected_comma=True))
    return findings


def ingest_md(text):
    """A table takes precedence; otherwise sectioned findings. (Deterministic rule
    to avoid double-counting a report that carries both a summary table and detail
    sections for the same findings.)"""
    findings = parse_md_table(text)
    if not findings:
        findings = parse_md_sections(text)
    engagement = {}
    m = re.search(r"^\s*#\s+(.+)$", text, re.M)
    if m:
        engagement["title"] = m.group(1).strip()
    return engagement, findings, []


# --- plaintext (pdf / docx fallback) ------------------------------------------
_PLAIN_HEAD_RE = re.compile(r"^\s*([A-Z]{1,6}-\d{1,4})\b[ \t]*[—–:.\)\-]?[ \t]*(.*)$")


def parse_plaintext(text):
    """Best-effort finding extraction from flat text (pdftotext / stripped docx):
    a `F-001 ...` id at line start opens a block; Severity:/CVSS:/CWE: lines fill it."""
    lines = text.splitlines()
    heads = []
    for i, ln in enumerate(lines):
        m = _PLAIN_HEAD_RE.match(ln)
        if m:
            heads.append((i, m.group(1), m.group(2).strip()))
    findings = []
    for idx, (li, fid, title) in enumerate(heads):
        end = heads[idx + 1][0] if idx + 1 < len(heads) else len(lines)
        values = {"id": fid, "title": title}
        for bl in lines[li + 1:end]:
            m = _FIELD_LINE_RE.match(bl)
            if m:
                key = _canon_header(m.group(1))
                if key:
                    values.setdefault(key, m.group(2).strip())
        if "severity" in values or title:
            findings.append(_build_finding(values, affected_comma=True))
    return findings


# --- pdf ----------------------------------------------------------------------
def _pdf_text(path):
    exe = shutil.which("pdftotext")
    if exe:
        try:
            out = subprocess.run([exe, "-layout", path, "-"],
                                 capture_output=True, text=True, timeout=180)
            if out.returncode == 0:
                return out.stdout, None
        except Exception as e:
            return None, "pdftotext failed: %s" % e
    try:
        from pdfminer.high_level import extract_text
        return extract_text(path), None
    except ImportError:
        return None, None
    except Exception as e:
        return None, "pdfminer failed: %s" % e


def ingest_pdf(path):
    text, err = _pdf_text(path)
    if text is None:
        msg = ("no PDF text extractor available "
               "(install poppler's pdftotext or pip install pdfminer.six)")
        raise MissingDependency(err + "; " + msg if err else msg)
    warnings = [err] if err else []
    findings = parse_plaintext(text)
    if not findings:
        warnings.append("no findings extracted from PDF text (best-effort)")
    return {}, findings, warnings


# --- docx ---------------------------------------------------------------------
def _strip_docx_xml(xml):
    xml = re.sub(r"</w:p>", "\n", xml)
    xml = re.sub(r"</w:tr>", "\n", xml)
    xml = re.sub(r"<w:tab[^>]*/>", "\t", xml)
    xml = re.sub(r"<[^>]+>", "", xml)
    return html.unescape(xml)


def ingest_docx(path):
    warnings = []
    tables_findings, text = [], None
    try:
        import docx
        doc = docx.Document(path)
        for t in doc.tables:
            grid = [[c.text for c in r.cells] for r in t.rows]
            fs, _ = _findings_from_grid(grid, affected_comma=True)
            tables_findings.extend(fs)
        text = "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        try:
            with zipfile.ZipFile(path) as z:
                text = _strip_docx_xml(z.read("word/document.xml").decode("utf-8", "replace"))
            warnings.append("python-docx not installed; used raw document.xml text")
        except Exception as e:
            raise MissingDependency(
                "python-docx not installed and raw docx read failed: %s" % e)
    except Exception as e:
        try:
            with zipfile.ZipFile(path) as z:
                text = _strip_docx_xml(z.read("word/document.xml").decode("utf-8", "replace"))
            warnings.append("python-docx read failed (%s); used raw document.xml text" % e)
        except Exception as e2:
            raise UnreadableSource("cannot read docx: %s" % e2)
    findings = tables_findings
    if not findings and text:
        findings = parse_md_table(text) or parse_md_sections(text) or parse_plaintext(text)
    if not findings:
        warnings.append("no findings extracted from DOCX (best-effort)")
    return {}, findings, warnings


# --- detection / orchestration ------------------------------------------------
_EXT_FORMAT = {".json": "json", ".xlsx": "xlsx", ".xlsm": "xlsx", ".md": "md",
               ".markdown": "md", ".pdf": "pdf", ".docx": "docx"}


def detect_format(path):
    ext = os.path.splitext(path)[1].lower()
    if ext in _EXT_FORMAT:
        return _EXT_FORMAT[ext]
    try:
        with open(path, "rb") as fh:
            head = fh.read(8)
    except OSError:
        return None
    if head[:4] == b"%PDF":
        return "pdf"
    if head[:4] == b"PK\x03\x04":
        try:
            with zipfile.ZipFile(path) as z:
                names = z.namelist()
            if any(n.startswith("word/") for n in names):
                return "docx"
            if any(n.startswith("xl/") for n in names):
                return "xlsx"
        except Exception:
            pass
        return "xlsx"
    try:
        with open(path, encoding="utf-8") as fh:
            s = fh.read(4096).lstrip()
        return "json" if s[:1] in "{[" else "md"
    except (OSError, UnicodeDecodeError):
        return None


_INGESTORS = {"xlsx": ingest_xlsx, "pdf": ingest_pdf, "docx": ingest_docx}


def build_result(source, fmt, engagement, findings, warnings):
    """Assemble the report_data-shaped result and run the non-null guards."""
    for f in findings:
        fid = f.get("id") or "?"
        if f.get("cwe") is None:
            warnings.append("%s: cwe is missing (fill before downstream use)" % fid)
        if f.get("cvss_vector") is None:
            warnings.append("%s: cvss_vector is missing (fill before downstream use)" % fid)
        if not f.get("severity"):
            warnings.append("%s: severity is missing" % fid)
        if not f.get("title"):
            warnings.append("%s: title is missing" % fid)
    return {"engagement": engagement or {}, "findings": findings,
            "ingest": {"source": source, "format": fmt, "warnings": warnings}}


def ingest(source, fmt="auto"):
    """Ingest `source` -> (result_dict_or_None, exit_code). result is None only on
    a usage/unreadable-source error (exit 2)."""
    if not os.path.isfile(source):
        return None, EXIT_USAGE
    if fmt in (None, "auto"):
        fmt = detect_format(source)
    if fmt is None:
        return None, EXIT_USAGE
    exit_code = EXIT_OK
    warnings = []
    engagement, findings = {}, []
    try:
        if fmt == "json":
            engagement, findings, w = ingest_json(source)
        elif fmt == "md":
            with open(source, encoding="utf-8", errors="replace") as fh:
                engagement, findings, w = ingest_md(fh.read())
        elif fmt in _INGESTORS:
            engagement, findings, w = _INGESTORS[fmt](source)
        else:
            return None, EXIT_USAGE
        warnings.extend(w)
    except MissingDependency as e:
        warnings.append(str(e))
        exit_code = EXIT_MISSING_DEP
    except UnreadableSource:
        return None, EXIT_USAGE
    return build_result(source, fmt, engagement, findings, warnings), exit_code


def main(argv=None):
    ap = argparse.ArgumentParser(
        description="Parse a finished/foreign report into canonical report_data JSON")
    ap.add_argument("source", help="report file (json/xlsx/md/pdf/docx)")
    ap.add_argument("--format", choices=["auto", "json", "xlsx", "md", "pdf", "docx"],
                    default="auto", help="force the source format (default: auto)")
    ap.add_argument("-o", "--output", help="write JSON here (default: stdout)")
    args = ap.parse_args(argv)

    result, code = ingest(args.source, args.format)
    if result is None:
        sys.stderr.write(
            "report_ingest: cannot ingest %r (unreadable or unknown format)\n"
            % args.source)
        return EXIT_USAGE
    payload = json.dumps(result, indent=2, ensure_ascii=False)
    if args.output:
        with open(args.output, "w", encoding="utf-8") as fh:
            fh.write(payload + "\n")
    else:
        sys.stdout.write(payload + "\n")
    for w in result["ingest"]["warnings"]:
        sys.stderr.write("warning: %s\n" % w)
    return code


if __name__ == "__main__":
    sys.exit(main())
