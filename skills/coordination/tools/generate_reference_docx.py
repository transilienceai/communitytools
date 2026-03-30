#!/usr/bin/env python3
"""
Generate a styled reference.docx for pandoc pentest report conversion.

Usage:
    python3 generate_reference_docx.py                          # Generate reference.docx
    python3 generate_reference_docx.py --post-process FILE.docx # Post-process a generated report

Requires: pip install python-docx
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# --- Color palette ---
NAVY = RGBColor(0x1B, 0x3A, 0x5C)
MEDIUM_BLUE = RGBColor(0x4A, 0x6F, 0xA5)
HEADING_BLUE = RGBColor(0x2E, 0x6C, 0xA8)
DARK_GRAY = RGBColor(0x3D, 0x3D, 0x3D)
BODY_GRAY = RGBColor(0x33, 0x33, 0x33)
BORDER_GRAY = RGBColor(0xD0, 0xD0, 0xD0)
ALT_ROW_BG = "F2F6FA"
TABLE_HEADER_BG = "1B3A5C"
CODE_BG = "F5F5F5"

SEV_CRITICAL = RGBColor(0xC0, 0x00, 0x00)
SEV_HIGH = RGBColor(0xED, 0x7D, 0x31)
SEV_MEDIUM = RGBColor(0xBF, 0x8F, 0x00)
SEV_LOW = RGBColor(0x54, 0x82, 0x35)
SEV_INFO = RGBColor(0x80, 0x80, 0x80)


def set_paragraph_spacing(style, before_pt=0, after_pt=0, line_spacing=None):
    """Set paragraph spacing on a style's paragraph format."""
    pf = style.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line_spacing:
        pf.line_spacing = line_spacing


def add_bottom_border(style, color_hex="1B3A5C", width=6):
    """Add a bottom border to a paragraph style."""
    pPr = style.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{width}" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def create_character_style(doc, name, color, bold=True, font_name="Calibri", size_pt=None):
    """Create a character style for inline severity coloring."""
    style = doc.styles.add_style(name, 2)  # WD_STYLE_TYPE.CHARACTER = 2
    style.font.color.rgb = color
    style.font.bold = bold
    style.font.name = font_name
    if size_pt:
        style.font.size = Pt(size_pt)
    return style


def configure_table_style(doc):
    """Create a professional table style with navy headers and alternating rows."""
    # python-docx doesn't fully support custom table styles,
    # so we configure defaults via Normal Table and rely on post-processing
    # for individual tables. The key styling happens in post_process_report().
    pass


def set_page_setup(doc):
    """Configure page margins, headers, and footers."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Header: CONFIDENTIAL
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("CONFIDENTIAL")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = "Calibri"

        # Footer: page numbers
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.name = "Calibri"
        # Page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._element.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._element.append(fldChar2)


def generate_reference_docx(output_path: str):
    """Generate the styled reference.docx template."""
    doc = Document()

    # --- Title style ---
    title_style = doc.styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = NAVY
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title_style, before_pt=0, after_pt=48)

    # --- Subtitle style ---
    subtitle_style = doc.styles["Subtitle"]
    subtitle_style.font.name = "Calibri"
    subtitle_style.font.size = Pt(16)
    subtitle_style.font.bold = False
    subtitle_style.font.color.rgb = MEDIUM_BLUE
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle_style, before_pt=0, after_pt=24)

    # --- Heading 1 ---
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    set_paragraph_spacing(h1, before_pt=24, after_pt=12)
    add_bottom_border(h1, "1B3A5C", width=6)

    # --- Heading 2 ---
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = HEADING_BLUE
    set_paragraph_spacing(h2, before_pt=18, after_pt=8)

    # --- Heading 3 ---
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = DARK_GRAY
    set_paragraph_spacing(h3, before_pt=12, after_pt=6)

    # --- Normal (body) ---
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_GRAY
    set_paragraph_spacing(normal, after_pt=6, line_spacing=1.15)

    # --- Block Text / Quote style (for code blocks) ---
    # Pandoc uses "Source Code" style for fenced code blocks
    try:
        source_code = doc.styles.add_style("Source Code", 1)  # PARAGRAPH
    except ValueError:
        source_code = doc.styles["Source Code"]
    source_code.font.name = "Consolas"
    source_code.font.size = Pt(10)
    source_code.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    set_paragraph_spacing(source_code, before_pt=4, after_pt=4)
    # Background shading for code blocks
    pPr = source_code.element.get_or_add_pPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>'
    )
    pPr.append(shading)

    # Also style "Verbatim Char" for inline code (pandoc uses this)
    try:
        verbatim = doc.styles.add_style("Verbatim Char", 2)  # CHARACTER
    except ValueError:
        verbatim = doc.styles["Verbatim Char"]
    verbatim.font.name = "Consolas"
    verbatim.font.size = Pt(10)
    verbatim.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Severity character styles ---
    create_character_style(doc, "Severity-Critical", SEV_CRITICAL)
    create_character_style(doc, "Severity-High", SEV_HIGH)
    create_character_style(doc, "Severity-Medium", SEV_MEDIUM)
    create_character_style(doc, "Severity-Low", SEV_LOW)
    create_character_style(doc, "Severity-Info", SEV_INFO)

    # --- TOC Heading ---
    try:
        toc_heading = doc.styles["TOC Heading"]
        toc_heading.font.name = "Calibri"
        toc_heading.font.size = Pt(18)
        toc_heading.font.bold = True
        toc_heading.font.color.rgb = NAVY
    except KeyError:
        pass

    # --- List styles ---
    for list_style_name in ["List Bullet", "List Number"]:
        try:
            ls = doc.styles[list_style_name]
            ls.font.name = "Calibri"
            ls.font.size = Pt(11)
            ls.font.color.rgb = BODY_GRAY
        except KeyError:
            pass

    # --- Page setup ---
    set_page_setup(doc)

    # --- Add a placeholder paragraph (pandoc needs content to read styles) ---
    p = doc.add_paragraph("Reference document for pandoc report generation.", style="Normal")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.save(output_path)
    print(f"Generated: {output_path}")


def post_process_report(docx_path: str):
    """
    Post-process a pandoc-generated DOCX to:
    1. Color severity text in table cells and paragraphs
    2. Style table headers with navy background and white text
    3. Apply alternating row shading
    4. Normalize column widths for common table patterns
    """
    doc = Document(docx_path)

    severity_map = {
        "CRITICAL": SEV_CRITICAL,
        "Critical": SEV_CRITICAL,
        "HIGH": SEV_HIGH,
        "High": SEV_HIGH,
        "MEDIUM": SEV_MEDIUM,
        "Medium": SEV_MEDIUM,
        "LOW": SEV_LOW,
        "Low": SEV_LOW,
        "INFORMATIONAL": SEV_INFO,
        "Informational": SEV_INFO,
        "INFO": SEV_INFO,
        "Info": SEV_INFO,
    }

    severity_pattern = re.compile(
        r'\b(CRITICAL|Critical|HIGH|High|MEDIUM|Medium|LOW|Low|INFORMATIONAL|Informational|INFO|Info)\b'
    )

    def color_severity_in_runs(paragraph):
        """Color severity keywords in paragraph runs."""
        for run in paragraph.runs:
            text = run.text
            if severity_pattern.search(text):
                for keyword, color in severity_map.items():
                    if keyword in text:
                        run.font.color.rgb = color
                        run.font.bold = True
                        break

    def style_table(table):
        """Apply professional styling to a table."""
        # Style header row
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # Navy background
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}" w:val="clear"/>'
                )
                tcPr.append(shading)
                # White bold text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)

        # Alternating row shading and severity coloring
        for i, row in enumerate(table.rows):
            if i == 0:
                continue  # skip header
            if i % 2 == 0:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{ALT_ROW_BG}" w:val="clear"/>'
                    )
                    tcPr.append(shading)

            # Color severity text in cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    color_severity_in_runs(paragraph)
                    # Style cell text
                    for run in paragraph.runs:
                        if run.font.color.rgb is None or run.font.color.rgb == RGBColor(0, 0, 0):
                            run.font.color.rgb = BODY_GRAY
                        run.font.name = "Calibri"
                        if run.font.size is None:
                            run.font.size = Pt(10)

        # Apply thin borders
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D0D0D0"/>'
            f'</w:tblBorders>'
        )
        # Remove existing borders if any
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(borders)

        # Set table width to 100%
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
            tblPr.append(tblW)
        else:
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')

    # Process all tables
    for table in doc.tables:
        style_table(table)

    # Color severity text in body paragraphs
    for paragraph in doc.paragraphs:
        color_severity_in_runs(paragraph)

    doc.save(docx_path)
    print(f"Post-processed: {docx_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate or post-process pentest report DOCX")
    parser.add_argument(
        "--post-process",
        metavar="FILE",
        help="Post-process an existing DOCX (color severity, style tables)",
    )
    parser.add_argument(
        "--output", "-o",
        default=str(Path(__file__).parent / "reference.docx"),
        help="Output path for reference.docx (default: same directory as script)",
    )
    args = parser.parse_args()

    if args.post_process:
        post_process_report(args.post_process)
    else:
        generate_reference_docx(args.output)


if __name__ == "__main__":
    main()
